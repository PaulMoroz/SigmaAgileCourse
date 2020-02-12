from PIL import Image
import sys
from fpdf import FPDF, HTMLMixin
from docx import Document
import pytesseract as tess
from table import Table
from table_main import find_table
import cv2 as cv
import numpy as np
import os


#=============================
#RECOGNIZING TEXT ON THE IMAGE
#=============================
def transform(image,to):
	if to == "PIL":
		img = cv.cvtColor(image, cv.COLOR_BGR2RGB)
		im_pil = Image.fromarray(img)
		return im_pil
	elif to=="cv2":
		opencvImage = cv.cvtColor(np.array(image), cv.COLOR_RGB2BGR)
		return opencvImage


def recognize(image):
	print("Recognizing",image)
	return tess.image_to_string(image)
 


#======================================
#Splitting images to table and no-table
#======================================
class HTML2PDF(FPDF, HTMLMixin):
    pass


#=======================
#Crop image
#=======================
def crop(image):
	cropped = []
	width,height = image.size
	tabs = find_table(transform(image,"cv2"))
	tabs.insert(0,Table(0,0,0,0)) 
	tabs.append(Table(0,height,0,0))
	for i in range(0,len(tabs)-1):
		cropped.append([image.crop((tabs[i].x,tabs[i].y,tabs[i].x+tabs[i].w,tabs[i].y+ tabs[i].h)),tabs[i]])
		cropped.append([image.crop((0,tabs[i].y+tabs[i].h,width,tabs[i+1].y)),"common"])
	cropped.append('eof')
	cropped = cropped[1:]
	print("Cropped file",len(cropped))
	return cropped 


#=======================
#Crop some images
#=======================


def image_split(image_path):
	cropped_all = []
	if sys.argv[4]=='file':
		if image_path.endswith("png") or image_path.endswith("jpg"):
			cropped_all.extend(crop(Image.open(image_path)))
		else:
			print("Error!")
	else:
		for file in os.listdir(image_path):
			if file.endswith("png") or file.endswith("jpg"):
				print("Enter file",file)
				p = image_path+'/'+file
				cropped_all.extend(crop(Image.open(image_path+'/'+file)))
	print("Len cropped all", cropped_all)
	return cropped_all


#==========================================
#CREATING PDF DOCUMENT WITH RECOGNIZED TEXT
#==========================================
def to_pdf():
	htmlcode = ""
	im_pieces = image_split(sys.argv[1])
	pdf = HTML2PDF()
	index = 1
	for k in range(len(im_pieces)):
		if im_pieces[k] == 'eof':
			if sys.argv[5] == 'm' or k == len(im_pieces)-1:
				pdf = HTML2PDF()
				pdf.add_page()
				pdf.write_html(htmlcode,'F')
				pdf.output(str(sys.argv[2])+(str(index) if sys.argv[5] == 'm' else "")+".pdf")
				index+=1
				htmlcode = ""
		elif im_pieces[k][1] == "common":
			htmlcode += "<p>" + str(recognize(im_pieces[k][0])) + "</p>"
		else:
			table_entries = im_pieces[k][1].get_table_entries()
			htmlcode += "<table>"
			table_roi = transform(im_pieces[k][0],"cv2")
			table_roi = cv.resize(table_roi, (im_pieces[k][0].w * 3, im_pieces[k][0].h * 3))
			try:
				for i in range(len(table_entries)):
					row = table_entries[i]
					htmlcode += "<tr>"
					for j in range(len(row)):  
						entry = row[j]
						entry_roi = table_roi[entry[1] * 3: (entry[1] + entry[3]) * 3, entry[0] * 3:(entry[0] + entry[2]) * 3]
						htmlcode += "<td " + "width = "+str(100//row)+"% > " + str(recognize(transform(table_roi[entry[1] * 3: (entry[1] + entry[3]) * 3, entry[0] * 3:(entry[0] + entry[2]) * 3],"PIL"))) + "</td>"
					htmlcode+="</tr>"
			except:
				print("Table Entries" ,len(table_entries))
			htmlcode+="</table>"



#===========================================
#CREATING WORD DOCUMENT WITH RECOGNIZED TEXT
#===========================================
def to_word():
	doc = Document()
	im_pieces,tabs = image_split(sys.argv[1]) 
	print("Lenght of image pieces",len(im_pieces))
	index = 1
	for k in range(t):
		if im_pieces[k] == 'eof':
			if sys.argv[5] == 'm' or k == len(im_pieces)-1:
			    doc.save(str(sys.argv[2])+str(index)+'.docx' if sys.argv[5]=='m' else str(sys.argv[2])+".docx")
			    doc = Document()
			    index+=1
		elif im_pieces[k][1] == "common" :
			doc.add_paragraph(recognize(im_pieces[k][0]))
		else:
			table_entries = im_pieces[k][1].get_table_entries()
			print("Rows, cols",len(table_entries),len(table_entries[0]))
			table = doc.add_table(rows = len(table_entries),cols =  len(table_entries[0]))
			table_roi = transform(im_pieces[k][0],"cv2")
			table_roi = cv.resize(table_roi, (im_pieces[k][0].w * 3, im_pieces[k][0].h * 3))
			try:
				for i in range(len(table_entries)):
					row = table_entries[i]
					for j in range(len(row)):  
						entry = row[j]
						entry_roi = table_roi[entry[1] * 3: (entry[1] + entry[3]) * 3, entry[0] * 3:(entry[0] + entry[2]) * 3]
						table.cell(i,j).text = str(recognize(transform(table_roi[entry[1] * 3: (entry[1] + entry[3]) * 3, entry[0] * 3:(entry[0] + entry[2]) * 3],"PIL")))
			except:
				print("Table Entries" ,len(table_entries))	



#==========================================
#CREATING TXT DOCUMENT WITH RECOGNIZED TEXT
#==========================================
def to_txt():
	im_pieces,tabs = image_split(sys.argv[1])
	t,index = len(im_pieces),1
	f = open(str(sys.argv[2])+(str(index) if sys.argv[5] == 'm' else "")+".txt","w+")
	for k in range(0,len(tabs)):
		if im_pieces[k] == 'eof':
			if sys.argv[5] == 'm':
				f.close()
				index+=1
				f = open(str(sys.argv[2])+str(index)+".txt","w+")
			elif sys.argv[k] == len(im_pieces) - 1:
				f.close()
		if im_pieces[k][1] == "common":
			f.write(recognize(im_pieces[k][0]))
		else:
			table_entries = im_pieces[k][1].get_table_entries()
			table_roi = transform(im_pieces[k][0],"cv2")
			table_roi = cv.resize(table_roi, (im_pieces[k][0].w * 3, im_pieces[k][0].h * 3))
			try:
				for i in range(len(table_entries)):
					row = table_entries[i]
					for j in range(len(row)):  
						entry = row[j]
						entry_roi = table_roi[entry[1] * 3: (entry[1] + entry[3]) * 3, entry[0] * 3:(entry[0] + entry[2]) * 3]
						f.write( str(recognize(transform(table_roi[entry[1] * 3: (entry[1] + entry[3]) * 3, entry[0] * 3:(entry[0] + entry[2]) * 3],"PIL"))) + "\t")
					f.write("\n")
			except:
				print("Table Entries" ,len(table_entries))



if __name__ == "__main__":
	if sys.argv[3] == "pdf":
		to_pdf()
	elif sys.argv[3] == "docx":
		to_word()
	else:
		to_txt()

